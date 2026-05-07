from __future__ import annotations

from docx import Document

from app.math_docx import add_math_aware_paragraph


def main() -> None:
    document = Document()
    document.add_heading("Kiểm tra Word Equation", level=1)
    paragraph = document.add_paragraph()
    add_math_aware_paragraph(
        paragraph,
        r"Ví dụ trong dòng: phương trình \(x^2-2x+1=0\), biệt thức \(\Delta=b^2-4ac\).",
    )
    paragraph = document.add_paragraph()
    add_math_aware_paragraph(paragraph, r"\[\frac{x^2+1}{x-1}=\sqrt{x+2}\]")
    output_path = "outputs/generated/math_equation_test.docx"
    document.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
