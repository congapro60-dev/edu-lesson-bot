from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document

from app.config import BASE_DIR
from app.drive_client import GoogleDriveClient


TEMPLATE_DIR = BASE_DIR / "outputs" / "templates"
SAMPLE_DOC_ID = "1N7xMK_heBVWRFU6aiL8aUGM-Us0NFsOs"
SAMPLE_DOCX_PATH = TEMPLATE_DIR / "detailed_sample_lesson.docx"


def extract_docx_text(path: Path, max_paragraphs: int = 120) -> list[str]:
    document = Document(path)
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
        if len(paragraphs) >= max_paragraphs:
            break
    return paragraphs


def extract_tables_text(path: Path, max_tables: int = 8) -> list[str]:
    document = Document(path)
    table_texts: list[str] = []
    for table_index, table in enumerate(document.tables[:max_tables], start=1):
        table_texts.append(f"--- Table {table_index} ---")
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells]
            if any(cells):
                table_texts.append(" || ".join(cells))
    return table_texts


def download_sample_doc() -> Path:
    TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
    client = GoogleDriveClient()
    metadata = client.get_file(SAMPLE_DOC_ID)
    mime_type = metadata.get("mimeType", "")
    if mime_type == "application/vnd.google-apps.document":
        return client.export_google_doc(SAMPLE_DOC_ID, SAMPLE_DOCX_PATH)
    return client.download_file(SAMPLE_DOC_ID, SAMPLE_DOCX_PATH)


def inspect_templates() -> None:
    files = [
        TEMPLATE_DIR / "huong_dan_soan_khdh_tds.docx",
        TEMPLATE_DIR / "mau_giao_an_ban_toan.docx",
        SAMPLE_DOCX_PATH,
    ]
    for file_path in files:
        print(f"\n--- {file_path.name} ---")
        if not file_path.exists():
            print(f"Missing file: {file_path}")
            continue
        paragraphs = extract_docx_text(file_path)
        for index, text in enumerate(paragraphs, start=1):
            print(f"{index:02d}. {text}")
        table_texts = extract_tables_text(file_path)
        for text in table_texts:
            print(text)


def main() -> None:
    parser = argparse.ArgumentParser(description="Inspect lesson template DOCX files")
    parser.add_argument("--download-sample", action="store_true", help="Download/export the detailed sample Google Doc")
    parser.add_argument("--inspect", action="store_true", help="Print non-empty paragraphs and tables from downloaded templates")
    args = parser.parse_args()

    if args.download_sample:
        output_path = download_sample_doc()
        print(f"Downloaded sample lesson: {output_path}")
    if args.inspect:
        inspect_templates()
    if not args.download_sample and not args.inspect:
        parser.print_help()


if __name__ == "__main__":
    main()
