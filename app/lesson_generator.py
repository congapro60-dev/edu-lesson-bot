from __future__ import annotations

import argparse
import re
import unicodedata
from dataclasses import dataclass, replace
from pathlib import Path

from app.drive_audit import PROGRAM_LABELS, audit_all_range, collect_missing_weeks, find_week_folder

from anthropic import Anthropic
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config import BASE_DIR, load_settings, require_values
from app.drive_client import GoogleDriveClient
from app.math_docx import add_math_aware_paragraph
from app.moet_parser import MoetWeekPlan, extract_moet_week
from app.pdf_renderer import render_pdf
from app.ppct_parser import TDS_EXCEL_PATH, LessonItem, TDSWeekPlan, extract_tds_week
from app.telegram_notify import build_notifier
from app.web_lesson_client import render_lesson_files_with_web, web_render_enabled


GENERATED_DIR = BASE_DIR / "outputs" / "generated"
TEMPLATE_DIR = BASE_DIR / "outputs" / "templates"
TDS_TEMPLATE_FILE_ID = "152TEJ_7P0RrqQbtk14TlQvPxJ8peNL2g"
TDS_TEMPLATE_DOCX_PATH = TEMPLATE_DIR / "mau_giao_an_tds.docx"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_MIME_TYPE = "application/pdf"


@dataclass(frozen=True)
class GeneratedLessonFiles:
    docx_path: Path
    pdf_path: Path | None = None
    lesson_title: str = ""
    uploaded_links: dict[str, str] | None = None


@dataclass(frozen=True)
class GeneratedLessonBatch:
    items: list[GeneratedLessonFiles]


@dataclass(frozen=True)
class ExistingOutputConflict:
    filename: str
    file_id: str
    web_view_link: str = ""
    mime_type: str = ""


# ---------------------------------------------------------------------------
# Prompt constants — đồng bộ với useLessonCreator.ts trên web
# ---------------------------------------------------------------------------

LESSON_SYSTEM_PROMPT = """Bạn là một CHUYÊN GIA GIÁO DỤC CAO CẤP với 20 năm kinh nghiệm thiết kế chương trình dạy học theo chuẩn quốc tế.
Nhiệm vụ: Soạn giáo án \"Masterpiece\" theo chuẩn WALT/WILF + Danielson Framework.
Phải tạo ra sản phẩm có độ chi tiết tối đa — minute-by-minute, không rút gọn, không generic.

Nguyên tắc bắt buộc:
1. Bám tuyệt đối PPCT được cung cấp; không bỏa bài học hay tài liệu.
2. Trả về toàn bộ nội dung trong thẻ <lesson_content>...</lesson_content>.
3. LaTeX: inline $...$, display $$...$$. TUYỆT ĐỐI KHÔNG dùng ký tự | trong công thức (dùng \\mid thay thế).
4. Mỗi lượt trao đổi GV↔HS = 1 hàng riêng trong bảng."""


_CLAUDE_FORMAT = """
===== MẪu GIÁO ÁN (BẮT BUỘC TUÂN THỦ TUYỆT ĐỐI) =====

BỐ CỤC: 5 HOẠT ĐỘNG + ĐÁNH GIÁ DANIELSON

# \U0001f4d8 GIÁO ÁN: [TÊN BÀI HỎc IN HOA]
**Môn:** Toán | **Lớp:** [lớp] | **Tuần:** [tuần] | **Thời lượng:** 40 phút

---

## \U0001f3af THÔNG TIN CHUNG

**WALT (We Are Learning To):**
> [Mục tiêu học tập 1-2 câu, giọng “chúng ta sẽ học cách...”]

**WILF (What I’m Looking For):**

| Mức độ | Yêu cầu |
|---|---|
| \U0001f336️ Cơ bản | [Yêu cầu tối thiểu — 2-3 ý cụ thể có công thức/ví dụ] |
| \U0001f336️\U0001f336️ Nâng cao | [Vận dụng linh hoạt — 2-3 dạng bài tiêu biểu] |
| \U0001f336️\U0001f336️\U0001f336️ Thách thức | [Chứng minh/sáng tạo/kết nối liên môn — 1-2 bài hóc búa] |

**NĂNG LỰC CỐT LÕI:**
- \U0001f9e0 Tư duy toán học: [Mô tả cụ thể]
- \U0001f4d0 Mô hình hóa toán học: [Mô tả cụ thể]
- \U0001f4ac Giao tiếp toán học: [Mô tả cụ thể]
- \U0001f527 Sử dụng công cụ: [Liệt kê công cụ cụ thể]

---

## \U0001f680 HOẠT ĐỘNG 1: MỞ ĐẦU (~5 phút)

**Mục tiêu:** [Tạo hứng thú, kích hoạt kiến thức nền, đặt vấn đề CỤ THỂ cho bài mới]

| Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến |
|---|---|---|
| [5-8 lượt thoại, GV verbatim trong “...”] | [HS1, HS2, HS3 phản hồi cụ thể] | [Nội dung bảng/câu hỏi mở] |

---

## \U0001f4da HOẠT ĐỘNG 2: HÌNH THÀNH KIẾN THỨC MỚI (~15 phút)

**Mục tiêu:** [Xây dựng kiến thức cốt lõi, rút ra tính chất quan trọng]

| Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến |
|---|---|---|
| [5-8 lượt Scaffolding từ dễ đến khó, dùng [Quét Radar], [\U0001f4a1 Tuyên ngôn: ...]] | [HS khám phá, phát biểu quy luật] | [Định lý/công thức LaTeX chính xác] |

---

## ✏️ HOẠT ĐỘNG 3: LUYỆN TậP (~10 phút)

**Mục tiêu:** Rèn kỹ năng — phân hóa 3 mức.

| Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến |
|---|---|---|
| **\U0001f336️ Bài 1 (Cơ bản):** \"[Đề bài cụ thể]\" | [Tự làm] | Bài 1: [Lời giải] ✅ |
| **\U0001f336️\U0001f336️ Bài 2 (Nâng cao):** \"[Đề bài cụ thể]\" | [Tự làm] | Bài 2: [Lời giải] ✅ |
| **\U0001f336️\U0001f336️\U0001f336️ Bài 3 (Thách thức):** \"[Đề bài]\" | [Tự làm] | Bài 3: [Lời giải] ∞ |

---

## \U0001f30d HOẠT ĐỘNG 4: VẬN DỤNG (~5 phút)

**Mục tiêu:** Liên hệ thực tế CỤ THỂ.

| Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến |
|---|---|---|
| [Tình huống thực tế cụ thể: y học/kinh tế/kỹ thuật/AI/môi trường] | [Tính toán/phân tích] | [Bài toán thực tế + kết quả LaTeX] |

---

## \U0001f4dd HOẠT ĐỘNG 5: SƠ KẾT — DẶN DÒ (~5 phút)

**Mục tiêu:** Hệ thống hóa, giao BTVN phân hóa.

| Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến |
|---|---|---|
| [Yêu cầu HS tóm tắt 3 ý chính] | [HS1, HS2, HS3 tóm tắt] | \U0001f4cb TÓM TẪT: 1️⃣ ... 2️⃣ ... 3️⃣ ... |
| [Giao BTVN phân hóa] | [Ghi BTVN] | \U0001f4cc BTVN: 1. Cơ bản 2. Nâng cao ⭐ Thách thức |

---

## \U0001f4cb Đánh giá của tổ trưởng chuyên môn

| Tiêu chí | Điểm | Nhận xét |
|---|---|---|
| 1a: Kiến thức chuyên môn & sư phạm | /4 | [Nhận xét cụ thể 2-3 câu] |
| 1b: Thấu hiểu học sinh | /4 | [Nhận xét cụ thể 2-3 câu] |
| 1c: Mục tiêu giảng dạy | /4 | [Nhận xét cụ thể 2-3 câu] |
| 1d: Tài nguyên dạy học | /4 | [Nhận xét cụ thể 2-3 câu] |
| 1e: Thiết kế bài giảng | /4 | [Nhận xét cụ thể 2-3 câu] |
| 1f: Đánh giá quá trình | /4 | [Nhận xét cụ thể 2-3 câu] |

**Tổng: /24**

QUY TẪC KHÔNG VI PHẠM:
1. Đủ 5 hoạt động — không gộp HOẠT ĐỘNG 4+5, không bỏ HOẠT ĐỘNG 5.
2. WILF đủ 3 mức \U0001f336️ / \U0001f336️\U0001f336️ / \U0001f336️\U0001f336️\U0001f336️.
3. Header bảng chính xác: \"Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng / Sản phẩm dự kiến\".
4. GV verbatim trong \"...\", không mô tả gián tiếp (\"GV nêu vấn đề\" — SAI).
5. HS phản hồi cụ thể (HS1, HS2...), không generic (\"HS trả lời\" — SAI).
6. 5-8 lượt thoại mỗi hoạt động.
7. BTVN có ít nhất 1 bài ⭐ (cho HS khá/giỏi).
8. Danielson: nhận xét 2-3 câu cụ thể mỗi tiêu chí (không generic).
9. Dew ey tuyên ngôn: lồng ghép [\U0001f4a1 Tuyên ngôn: ...] vào ít nhất 2 hoạt động.
===== KẾT THÚC MẪu =====
"""


_MATH_RESTRICTIONS = """
===========================================================
QUY TẪC MÔN TOÁN — BẮT BUỘC:
I. MỤC TIÊU: Tư duy toán học, Mô hình hóa, Giao tiếp toán học, Giải quyết vấn đề, Công cụ & phương tiện.
   Phân hóa: HS khá/giỏi (nâng cao cụ thể) + HS TB/yếu (tối thiểu cần đạt, hỗ trợ cụ thể).

II. BẢNG 3 CỘT bắt buộc cho TẤT CẢ hoạt động:
   | Hoạt động của GV | Hoạt động của HS | Nội dung ghi bảng/Sản phẩm dự kiến |
   Mỗi lượt GV↔HS = 1 hàng riêng. TUYỆT ĐỐI KHÔNG dùng <br/><br/> để gộp nhiều lượt vào 1 hàng.

III. LATEX: $...$ inline, $$...$$ display. TUYỆT ĐỐI KHÔNG dùng | trong công thức → dùng \\mid.
===========================================================
"""


def tds_grade_output_folder_id(grade: int) -> str:
    settings = load_settings()
    if grade == 10:
        return settings.tds_g10_folder_id
    if grade == 11:
        return settings.tds_g11_folder_id
    if grade == 12:
        return settings.tds_g12_folder_id
    raise ValueError(f"Unsupported grade: {grade}")


def moet_grade_output_folder_id(grade: int) -> str:
    settings = load_settings()
    if grade == 10:
        return settings.moet_g10_folder_id
    if grade == 11:
        return settings.moet_g11_folder_id
    if grade == 12:
        return settings.moet_g12_folder_id
    raise ValueError(f"Unsupported grade: {grade}")


def build_lesson_prompt(plan: TDSWeekPlan | MoetWeekPlan, program: str = "TDS") -> str:
    """Build the user-turn prompt for lesson generation, matching the web's structure."""
    lessons_text = "\n".join(
        f"  - Tiết {lesson.period}: {lesson.content}"
        for lesson in plan.lessons
    )
    week_label = getattr(plan, "week_label", f"Tuần {plan.week:02d}")
    month = getattr(plan, "month", "")
    notes = getattr(plan, "notes", "")
    month_part = f" Tháng: {month}." if month else ""
    notes_part = f" Ghi chú PPCT: {notes}." if notes else ""

    return (
        f"BẠN LÀ MỘT CHUYÊN GIA GIÁO DỤC CAO CẤP.\n"
        f"NHIỆM VỤ: Soạn một giáo án \"Masterpiece\" (Kiệt tác sư phạm).\n\n"
        f"BỐ CỤC PHẢN HỒI BẮT BUỘC:\n"
        f"1. <thinking>: Phân tích mục tiêu bài học, đặc điểm HS lớp {plan.grade}, lựa chọn phương pháp.\n"
        f"2. <lesson_content>: TOÀN BỘ giáo án chi tiết (Markdown), BAO GỒM cả đánh giá Danielson ở cuối.\n\n"
        f"THÔNG TIN BÀI HỌC TỪ PPCT:\n"
        f"- Chương trình: Toán {program}. Lớp: {plan.grade}. {week_label}.{month_part}{notes_part}\n"
        f"- Số tiết tuần này: {len(plan.lessons)}\n"
        f"- Nội dung PPCT:\n{lessons_text}\n"
        f"{_CLAUDE_FORMAT}\n"
        f"{_MATH_RESTRICTIONS}"
    )


def generate_lesson_text(plan: TDSWeekPlan | MoetWeekPlan, program: str = "TDS") -> str:
    settings = load_settings()
    require_values(settings, ["anthropic_api_key"])
    client = Anthropic(api_key=settings.anthropic_api_key)
    try:
        response = client.messages.create(
            model=settings.claude_model,
            max_tokens=16000,
            system=LESSON_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": build_lesson_prompt(plan, program)}],
        )
    except Exception as exc:
        raise RuntimeError(
            "Không tạo được giáo án. API lỗi. "
            f"Chi tiết: {exc}"
        ) from exc

    raw_text = "\n".join(
        block.text for block in response.content if getattr(block, "type", "") == "text"
    ).strip()

    # Trích xuất <lesson_content> — giống hàm extractLessonContent() trên web
    content_match = re.search(r"<lesson_content>([\s\S]*?)</lesson_content>", raw_text, re.IGNORECASE)
    if content_match:
        lesson_text = content_match.group(1).strip()
    else:
        lesson_text = re.sub(r"<thinking>[\s\S]*?</thinking>", "", raw_text, flags=re.IGNORECASE).strip()

    if not lesson_text:
        raise RuntimeError("API không trả về giáo án hợp lệ; bot dừng để tránh upload file rỗng.")
    return lesson_text


def clean_markdown_inline(text: str) -> str:
    cleaned = text.replace("<br/>", "\n").replace("<br />", "\n").replace("<br>", "\n")
    cleaned = re.sub(r"(?<!\\)\*\*(.+?)(?<!\\)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"(?<!\\)__(.+?)(?<!\\)__", r"\1", cleaned)
    cleaned = cleaned.replace("`", "")
    return cleaned.strip()


def set_run_font(paragraph, size: int = 12, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bool(run.bold or bold)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_multiline_paragraph(document: Document, text: str) -> None:
    render_markdown_to_docx(document, text)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: int = 12, alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    cell.text = ""
    lines = clean_markdown_inline(text).splitlines() or [""]
    first_paragraph = cell.paragraphs[0]
    add_math_aware_paragraph(first_paragraph, lines[0])
    for line in lines[1:]:
        paragraph = cell.add_paragraph()
        add_math_aware_paragraph(paragraph, line)
    for paragraph in cell.paragraphs:
        if alignment is not None:
            paragraph.alignment = alignment
        set_run_font(paragraph, font_size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table) -> None:
    table_element = table._tbl
    table_properties = table_element.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def apply_table_grid_style(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def set_cell_width(cell, width_twips: int) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    width = cell_properties.first_child_found_in("w:tcW")
    if width is None:
        width = OxmlElement("w:tcW")
        cell_properties.append(width)
    width.set(qn("w:w"), str(width_twips))
    width.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str = "E2E8F0") -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_activity_table_widths(table, col_count: int) -> None:
    if col_count != 3:
        return
    widths = [1700, 4200, 4100]
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=False):
            set_cell_width(cell, width)


def add_table_row(table, values: list[str], *, bold: bool = False) -> None:
    row = table.add_row()
    for index, value in enumerate(values):
        set_cell_text(row.cells[index], value, bold=bold, alignment=WD_ALIGN_PARAGRAPH.CENTER if bold else None)


def lesson_title(plan: TDSWeekPlan | MoetWeekPlan) -> str:
    if not plan.lessons:
        return "Kế hoạch dạy học theo PPCT"
    first = plan.lessons[0].content.strip()
    if len(plan.lessons) == 1:
        return first
    return f"{first} và các nội dung tuần {plan.week}"


def ppct_summary(plan: TDSWeekPlan | MoetWeekPlan) -> str:
    return "\n".join(f"Tiết {lesson.period}: {lesson.content}" for lesson in plan.lessons)


def safe_filename_part(value: str, fallback: str = "giao-an") -> str:
    normalized = unicodedata.normalize("NFC", value or "")
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', " ", normalized)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or fallback


def limit_filename_stem(stem: str, max_length: int = 150) -> str:
    if len(stem) <= max_length:
        return stem
    return stem[:max_length].rstrip(" .,-–—")


def lesson_filename_prefix(lesson: LessonItem, index: int) -> str:
    content = safe_filename_part(lesson.content, f"Bài {index:02d}")
    period = safe_filename_part(lesson.period, "")
    if period:
        return limit_filename_stem(f"Tiết {period} - {content}")
    return limit_filename_stem(content)


def single_lesson_plan(plan: TDSWeekPlan | MoetWeekPlan, lesson: LessonItem) -> TDSWeekPlan | MoetWeekPlan:
    if isinstance(plan, TDSWeekPlan):
        return TDSWeekPlan(
            grade=plan.grade,
            week=plan.week,
            week_label=plan.week_label,
            month=plan.month,
            track=plan.track,
            notes=plan.notes,
            lessons=[lesson],
        )
    return MoetWeekPlan(grade=plan.grade, week=plan.week, lessons=[lesson])


def expected_lesson_file_names(plan: TDSWeekPlan | MoetWeekPlan) -> list[str]:
    names: list[str] = []
    for index, lesson in enumerate(plan.lessons, start=1):
        prefix = lesson_filename_prefix(lesson, index)
        names.extend([f"{prefix}.docx", f"{prefix}.pdf"])
    return names


def load_week_plan(program: str, grade: int, week: int, track: str = "dgs") -> TDSWeekPlan | MoetWeekPlan:
    normalized = program.lower().strip()
    if normalized == "tds":
        return extract_tds_week(TDS_EXCEL_PATH, grade, week, track)
    if normalized == "moet":
        return extract_moet_week(grade, week)
    raise ValueError(f"Unsupported program: {program}")


def output_parent_folder_id(program: str, grade: int) -> str:
    normalized = program.lower().strip()
    if normalized == "tds":
        return tds_grade_output_folder_id(grade)
    if normalized == "moet":
        return moet_grade_output_folder_id(grade)
    raise ValueError(f"Unsupported program: {program}")


def find_existing_output_conflicts(
    program: str,
    grade: int,
    week: int,
    track: str = "dgs",
) -> list[ExistingOutputConflict]:
    plan = load_week_plan(program, grade, week, track)
    expected_names = set(expected_lesson_file_names(plan))
    if not expected_names:
        return []

    client = GoogleDriveClient()
    week_folder = find_week_folder(client, output_parent_folder_id(program, grade), week)
    if not week_folder:
        return []

    files = client.list_files(query=f"'{week_folder['id']}' in parents and trashed = false", page_size=100)
    conflicts: list[ExistingOutputConflict] = []
    for file in files:
        name = file.get("name", "")
        if name in expected_names:
            conflicts.append(
                ExistingOutputConflict(
                    filename=name,
                    file_id=file.get("id", ""),
                    web_view_link=file.get("webViewLink", ""),
                    mime_type=file.get("mimeType", ""),
                )
            )
    return conflicts


def apply_base_style(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            styles[style_name].font.bold = True
            styles[style_name].font.size = Pt(13 if style_name != "Heading 1" else 14)


def add_section_table(document: Document, title: str, instruction: str, body: str) -> None:
    table = document.add_table(rows=2, cols=1)
    apply_table_grid_style(table)
    set_table_borders(table)
    set_cell_text(table.rows[0].cells[0], f"{title}\n{instruction}".strip(), bold=True, font_size=12)
    shade_cell(table.rows[0].cells[0])
    set_cell_text(table.rows[1].cells[0], body, font_size=12)


def add_activity_table(document: Document, title: str, instruction: str, rows: list[list[str]]) -> None:
    header = document.add_table(rows=1, cols=1)
    apply_table_grid_style(header)
    set_table_borders(header)
    set_cell_text(header.rows[0].cells[0], f"{title}\n{instruction}".strip(), bold=True, font_size=12)
    shade_cell(header.rows[0].cells[0], "D9EAD3")

    table = document.add_table(rows=1, cols=3)
    apply_table_grid_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for index, value in enumerate(["Thời gian thực", "Giáo viên và Học sinh", "Nội dung"]):
        set_cell_text(table.rows[0].cells[index], value, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index])
    for row in rows:
        add_table_row(table, row)
    set_activity_table_widths(table, 3)


def is_markdown_table_separator(line: str) -> bool:
    cells = split_markdown_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def split_markdown_row(line: str) -> list[str]:
    stripped = line.strip()
    if not stripped.startswith("|") or "|" not in stripped[1:]:
        return []
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    stripped = stripped[1:]
    return [cell.strip() for cell in re.split(r"(?<!\\)\|", stripped)]


def add_markdown_table(document: Document, table_rows: list[list[str]]) -> None:
    rows = [row for row in table_rows if not (len(row) > 0 and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row))]
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    if col_count <= 0:
        return

    table = document.add_table(rows=1, cols=col_count)
    apply_table_grid_style(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    for index in range(col_count):
        value = rows[0][index] if index < len(rows[0]) else ""
        set_cell_text(table.rows[0].cells[index], value, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[index])

    for source_row in rows[1:]:
        row = table.add_row()
        for index in range(col_count):
            value = source_row[index] if index < len(source_row) else ""
            set_cell_text(row.cells[index], value)

    set_activity_table_widths(table, col_count)


def add_markdown_paragraph(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return
    if stripped.startswith("#"):
        level = min(max(len(stripped) - len(stripped.lstrip("#")), 1), 3)
        paragraph = document.add_heading("", level=level)
        add_math_aware_paragraph(paragraph, clean_markdown_inline(stripped.lstrip("#").strip()))
        set_run_font(paragraph, 13 if level > 1 else 14, True)
        return
    if stripped.startswith("-"):
        paragraph = document.add_paragraph()
        add_math_aware_paragraph(paragraph, clean_markdown_inline(f"• {stripped[1:].strip()}"))
        set_run_font(paragraph, 12)
        return
    numbered_match = re.match(r"^(\d+\.)\s+", stripped)
    if numbered_match:
        paragraph = document.add_paragraph()
        body = re.sub(r"^\d+\.\s+", "", stripped)
        add_math_aware_paragraph(paragraph, clean_markdown_inline(f"{numbered_match.group(1)} {body}"))
        set_run_font(paragraph, 12)
        return
    if stripped.startswith(">"):
        paragraph = document.add_paragraph()
        add_math_aware_paragraph(paragraph, clean_markdown_inline(stripped.lstrip(">").strip()))
        set_run_font(paragraph, 12, False)
        return

    paragraph = document.add_paragraph()
    add_math_aware_paragraph(paragraph, clean_markdown_inline(stripped))
    set_run_font(paragraph, 12)


def render_markdown_to_docx(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    position = 0
    while position < len(lines):
        line = lines[position]
        if split_markdown_row(line):
            table_rows: list[list[str]] = []
            while position < len(lines) and split_markdown_row(lines[position]):
                row = split_markdown_row(lines[position])
                table_rows.append(row)
                position += 1
            add_markdown_table(document, table_rows)
            continue
        add_markdown_paragraph(document, line)
        position += 1


def download_tds_template_if_missing() -> Path | None:
    if TDS_TEMPLATE_DOCX_PATH.exists():
        return TDS_TEMPLATE_DOCX_PATH
    try:
        TEMPLATE_DIR.mkdir(parents=True, exist_ok=True)
        client = GoogleDriveClient()
        metadata = client.get_file(TDS_TEMPLATE_FILE_ID)
        mime_type = metadata.get("mimeType", "")
        if mime_type == "application/vnd.google-apps.document":
            return client.export_google_doc(TDS_TEMPLATE_FILE_ID, TDS_TEMPLATE_DOCX_PATH)
        return client.download_file(TDS_TEMPLATE_FILE_ID, TDS_TEMPLATE_DOCX_PATH)
    except Exception as exc:
        print(f"Could not download TDS template, falling back to blank DOCX: {exc}")
        return None


def clear_document_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def create_document_from_template() -> Document:
    template_path = download_tds_template_if_missing()
    if template_path and template_path.exists():
        try:
            document = Document(str(template_path))
            clear_document_body(document)
            return document
        except Exception as exc:
            print(f"Could not open TDS template, falling back to blank DOCX: {exc}")
    return Document()


def build_docx(
    plan: TDSWeekPlan | MoetWeekPlan,
    lesson_text: str,
    program: str = "TDS",
    filename_prefix: str | None = None,
) -> Path:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    output_stem = filename_prefix or f"{program}_G{plan.grade}_Tuan_{plan.week:02d}_{safe_filename_part(lesson_title(plan))}"
    output_path = GENERATED_DIR / f"{output_stem}.docx"

    document = create_document_from_template()
    apply_base_style(document)

    notes = getattr(plan, "notes", "")
    title = lesson_title(plan)
    heading = document.add_heading("KẾ HOẠCH DẠY HỌC", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_table = document.add_table(rows=2, cols=6)
    apply_table_grid_style(info_table)
    set_table_borders(info_table)
    rows = [
        ["Lớp", str(plan.grade), "Tên bài học", title, "Môn học", "Toán"],
        ["Giáo viên", "TDS THT", "Tuần học", str(plan.week), "Năm học", "2025 – 2026"],
    ]
    for row, values in zip(info_table.rows, rows, strict=True):
        for index, value in enumerate(values):
            is_label = index % 2 == 0
            set_cell_text(row.cells[index], value, bold=is_label, alignment=WD_ALIGN_PARAGRAPH.CENTER if is_label else None)
            if is_label:
                shade_cell(row.cells[index])

    materials = "\n".join([
        f"Nội dung PPCT {program}:",
        ppct_summary(plan),
    ])
    if notes:
        materials += f"\nGhi chú PPCT: {notes}"
    add_section_table(document, "PPCT TUẦN", "Dữ liệu nguồn dùng để soạn giáo án", materials)

    render_markdown_to_docx(document, lesson_text)

    document.save(output_path)
    return output_path


def get_existing_or_create_week_folder(client: GoogleDriveClient, parent_id: str, week: int) -> dict[str, str]:
    existing_folder = find_week_folder(client, parent_id, week)
    if existing_folder:
        return existing_folder
    return client.get_or_create_child_folder(parent_id, f"Tuần {week}")


def upload_generated_files(
    files: GeneratedLessonFiles,
    parent_id: str,
    week: int,
    replace_existing: bool = False,
) -> dict[str, str]:
    client = GoogleDriveClient()
    week_folder = get_existing_or_create_week_folder(client, parent_id, week)
    links: dict[str, str] = {}

    uploaded_docx = client.upload_file(files.docx_path, week_folder["id"], DOCX_MIME_TYPE, replace_existing=replace_existing)
    links["docx"] = uploaded_docx.get("webViewLink", "")
    print(f"Uploaded DOCX: {uploaded_docx.get('name')} | {uploaded_docx.get('id')} | {links['docx']}")

    if files.pdf_path:
        uploaded_pdf = client.upload_file(files.pdf_path, week_folder["id"], PDF_MIME_TYPE, replace_existing=replace_existing)
        links["pdf"] = uploaded_pdf.get("webViewLink", "")
        print(f"Uploaded PDF: {uploaded_pdf.get('name')} | {uploaded_pdf.get('id')} | {links['pdf']}")

    return links


def render_single_lesson_files(
    plan: TDSWeekPlan | MoetWeekPlan,
    lesson: LessonItem,
    index: int,
    program: str,
) -> GeneratedLessonFiles:
    one_lesson_plan = single_lesson_plan(plan, lesson)
    title = lesson.content.strip() or f"Bài {index:02d}"
    filename_prefix = lesson_filename_prefix(lesson, index)
    lesson_text = generate_lesson_text(one_lesson_plan, program)

    if web_render_enabled():
        rendered_files = render_lesson_files_with_web(
            title=title,
            content=lesson_text,
            output_dir=GENERATED_DIR,
            grade=plan.grade,
            week=plan.week,
            program=program,
            lesson_name=title,
            filename_prefix=filename_prefix,
        )
        docx_path = rendered_files.docx_path
        pdf_path = rendered_files.pdf_path
    else:
        docx_path = build_docx(one_lesson_plan, lesson_text, program, filename_prefix)
        pdf_path = render_pdf(one_lesson_plan, lesson_text, program, filename_prefix)

    print(f"Generated standard lesson DOCX: {docx_path}")
    print(f"Generated standard lesson PDF: {pdf_path}")
    return GeneratedLessonFiles(docx_path=docx_path, pdf_path=pdf_path, lesson_title=title)



def generate_lesson_batch(
    plan: TDSWeekPlan | MoetWeekPlan,
    program: str,
    parent_id: str,
    upload: bool = False,
    notify: bool = False,
    replace_existing: bool = False,
) -> GeneratedLessonBatch:
    if not plan.lessons:
        raise RuntimeError(f"Không tìm thấy bài học PPCT cho {program} G{plan.grade} tuần {plan.week:02d}.")

    generated_items = [
        render_single_lesson_files(plan, lesson, index, program)
        for index, lesson in enumerate(plan.lessons, start=1)
    ]

    if upload:
        generated_items = [
            replace(files, uploaded_links=upload_generated_files(files, parent_id, plan.week, replace_existing))
            for files in generated_items
        ]

    batch = GeneratedLessonBatch(items=generated_items)
    if notify:
        message_lines = [
            f"Đã tạo giáo án chuẩn {program} G{plan.grade} tuần {plan.week:02d}: {len(batch.items)} bài riêng biệt.",
        ]
        for item in batch.items:
            message_lines.append(f"- {item.docx_path.name}")
            if item.uploaded_links:
                for label, link in item.uploaded_links.items():
                    if link:
                        message_lines.append(f"  {label.upper()}: {link}")
        build_notifier().send_message("\n".join(message_lines))
        print("Telegram notification sent.")

    return batch



def generate_tds_docx(
    grade: int,
    week: int,
    track: str,
    upload: bool = False,
    notify: bool = False,
    replace_existing: bool = False,
) -> GeneratedLessonBatch:
    plan = extract_tds_week(TDS_EXCEL_PATH, grade, week, track)
    return generate_lesson_batch(plan, "TDS", tds_grade_output_folder_id(grade), upload, notify, replace_existing)



def generate_moet_docx(
    grade: int,
    week: int,
    upload: bool = False,
    notify: bool = False,
    replace_existing: bool = False,
) -> GeneratedLessonBatch:
    plan = extract_moet_week(grade, week)
    return generate_lesson_batch(plan, "MOET", moet_grade_output_folder_id(grade), upload, notify, replace_existing)

def selected_programs(include_tds: bool, include_moet: bool) -> list[str]:
    programs: list[str] = []
    if include_tds:
        programs.append("tds")
    if include_moet:
        programs.append("moet")
    return programs


def generate_missing_docx(
    programs: list[str],
    start_week: int,
    end_week: int,
    upload: bool = True,
    notify: bool = True,
    generate_partial_weeks: bool = False,
) -> None:
    audit_all_range(programs, start_week, end_week, notify)
    missing = collect_missing_weeks(programs, start_week, end_week, generate_partial_weeks)
    if not missing:
        message = "Không phát hiện tuần trống hoàn toàn cần tạo trong phạm vi đã kiểm tra."
        print(message)
        if notify:
            build_notifier().send_message(message)
        return

    summary_lines = ["Bắt đầu tạo giáo án cho các tuần còn thiếu:"]
    for program, grade, week, audit in missing:
        summary_lines.append(
            f"- {PROGRAM_LABELS[program]} G{grade} tuần {week:02d}: còn thiếu ước lượng {audit.missing_count} file"
        )
    summary = "\n".join(summary_lines)
    print(summary)
    if notify:
        build_notifier().send_message(summary)

    for program, grade, week, _audit in missing:
        if program == "tds":
            generate_tds_docx(grade, week, "dgs", upload, notify)
        elif program == "moet":
            generate_moet_docx(grade, week, upload, notify)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate lesson-plan DOCX files")
    parser.add_argument("--tds", action="store_true", help="Generate one TDS lesson-plan DOCX")
    parser.add_argument("--moet", action="store_true", help="Generate one Moet lesson-plan DOCX")
    parser.add_argument("--missing", action="store_true", help="Audit first, then generate only missing weekly lesson-plan DOCX files")
    parser.add_argument("--start-week", type=int, default=1)
    parser.add_argument("--end-week", type=int, default=1)
    parser.add_argument("--grade", type=int, choices=[10, 11, 12], default=10)
    parser.add_argument("--week", type=int, default=1)
    parser.add_argument("--track", choices=["dgs", "discover"], default="dgs")
    parser.add_argument("--partial-weeks", action="store_true", help="Also regenerate weeks that already have some DOCX files but still look incomplete")
    parser.add_argument("--upload", action="store_true", help="Upload generated DOCX to Google Drive")
    parser.add_argument("--notify", action="store_true", help="Send Telegram notification after generation")
    args = parser.parse_args()

    if args.missing:
        programs = selected_programs(args.tds, args.moet) or ["tds", "moet"]
        generate_missing_docx(programs, args.start_week, args.end_week, args.upload, args.notify, args.partial_weeks)
    elif args.tds:
        generate_tds_docx(args.grade, args.week, args.track, args.upload, args.notify)
    elif args.moet:
        generate_moet_docx(args.grade, args.week, args.upload, args.notify)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
